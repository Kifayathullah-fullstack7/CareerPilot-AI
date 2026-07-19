import os
import pdfplumber
import docx

def extract_text_from_pdf(file_path):
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def extract_text_from_docx(file_path):
    text = []
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return "\n".join(text)

def parse_resume(file_path):
    if not os.path.exists(file_path):
        return ""
    
    ext = file_path.rsplit(".", 1)[1].lower() if "." in file_path else ""
    if ext == "pdf":
        return extract_text_from_pdf(file_path)
    elif ext == "docx" or ext == "doc":
        return extract_text_from_docx(file_path)
    return ""
